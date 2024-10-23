from langchain.chains import RetrievalQA
from langchain.vectorstores import Chroma
from langchain.embeddings import HuggingFaceInstructEmbeddings
from langchain.llms import HuggingFacePipeline
from langchain import PromptTemplate
from langchain.chains import ConversationChain, ConversationalRetrievalChain
from langchain.memory import VectorStoreRetrieverMemory
import torch
from constants import CHROMA_SETTINGS, SOURCE_DIRECTORY, PERSIST_DIRECTORY, LLM_MODEL, EMBEDDINGS_MODEL, TASK_NAME
from transformers import LlamaTokenizer, LlamaForCausalLM, pipeline
import click
import pyodbc
import time
from ingest import load_documents
from custom_memory import custom_memory
from json_extract import json_extract
import sys
import re 
import faiss
from langchain.docstore import InMemoryDocstore
from langchain.vectorstores import FAISS
from langchain.embeddings.openai import OpenAIEmbeddings
from transformers import AutoModelForCausalLM, AutoTokenizer, GenerationConfig
import openai
import openpyxl
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import docx
from pdf2docx import Converter
from os.path import exists


from constants import CHROMA_SETTINGS
from datetime import datetime


def load_model():
    """
    Select a model on huggingface.
    If you are running this for the first time, it will download a model for you.
    subsequent runs will use the model from the disk.
    """
    model_id = LLM_MODEL
    #tokenizer = AutoTokenizer.from_pretrained("tiiuae/falcon-40b-instruct")
    #model = AutoModelForCausalLM.from_pretrained("tiiuae/falcon-40b-instruct", trust_remote_code=True)
    #tokenizer = AutoTokenizer.from_pretrained("Open-Orca/OpenOrca-Platypus2-13B")
    #model = AutoModelForCausalLM.from_pretrained("Open-Orca/OpenOrca-Platypus2-13B")

    print(model_id)
    tokenizer = LlamaTokenizer.from_pretrained(model_id)

    model = LlamaForCausalLM.from_pretrained(model_id, #model_basename
                                             #   load_in_8bit=True, # set these options if your GPU supports them!
                                             #   device_map=1#'auto',
                                             #   torch_dtype=torch.float16,
                                             #   low_cpu_mem_usage=True
                                             )
    
    task_name = TASK_NAME

    pipe = pipeline(
        task_name,
        model=model,
        tokenizer=tokenizer,
        max_length=4000,
        temperature=0,
        top_p=0.95,
        repetition_penalty=1.15
    )

    local_llm = HuggingFacePipeline(pipeline=pipe)

    # tokenizer = AutoTokenizer.from_pretrained(model_id)

    # model = AutoModelForCausalLM.from_pretrained(
    #         model_id,
    #         device_map="auto",
    #         torch_dtype=torch.float32,
    #         low_cpu_mem_usage=True,
    #         trust_remote_code=True,
    #         # max_memory={0: "15GB"} # Uncomment this line with you encounter CUDA out of memory errors
    #     )
    # model.tie_weights()

    # generation_config = GenerationConfig.from_pretrained(model_id)

    # pipe = pipeline(
    #     "text-generation",
    #     model=model,
    #     tokenizer=tokenizer,
    #     max_length=2048,
    #     temperature=0,
    #     top_p=0.95,
    #     repetition_penalty=1.15,
    #     generation_config=generation_config,
    # )

    # local_llm = HuggingFacePipeline(pipeline=pipe)

    return local_llm

@click.command()
@click.option('--device_type', default='cuda', help='device to run on, select gpu, cpu or mps')
def main(device_type, ):

    device_type = "cuda"
    # load the instructorEmbeddings
    if device_type in ['cpu', 'CPU']:
        device='cpu'
    elif device_type in ['mps', 'MPS']:
        device='mps'
    else:
        device='cuda'

    print(f"Running on: {device}")

    print(torch.cuda.is_available())

    # Load embeddings
    embeddings = HuggingFaceInstructEmbeddings(model_name=EMBEDDINGS_MODEL,
                                               model_kwargs={"device": device})

    # Convert pdf to doc if needed
    #pdftodocx()

    # Extract section headers from doc
    documents = load_documents(SOURCE_DIRECTORY)
    sections = extract_headers(documents[0])
    print(sections)

    # Specialized filters on sections (Uncomment only if needed)
    #filtered_sections = [(section, re.sub(r'\s+\d+$', '', title.strip())) for section, title in sections if '(U)' in title]
    filtered_sections = [(section, re.sub(r'\s+\d+$', '', title.strip())) for section, title in sections]
    print(filtered_sections)

    # Put string tuples together and save them in new list
    section_strings = []
    for sec in filtered_sections:
        #string = sec[0] + ' ' + sec[1]
        string = sec[1]
        section_strings.append(string)
    print(section_strings)
    
    l, r = 0, 1
    final_sections = []

    # Iterates and extracts section content between two section titles
    while r < len(section_strings):
        extraction = match_and_extract(documents[0], section_strings[l], section_strings[r])
        if ' '.join(extraction) != '':
            final_sections.append(' '.join(extraction))
        l += 1
        r += 1
    print(final_sections)
    
    # Load the LLM for generating Natural Language responses. 
    llm = load_model()

    openai.api_key = ('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')

    
    # Template
    _DEFAULT_TEMPLATE = """The following is a conversation between a client and AI. The AI is acting as a product manager at NAI (North Atlantic Industries), a world leading independent supplier of rugged COTS embedded computing products for the industrial, commercial aerospace, and defense markets. The client 
    has sent NAI a Request for Proposal (RFP). You are in charge of using your memory of previous responses to sections to respond to each section of the Request for Proposal. If the client asks you to respond to a section you have responded to before, return to me the past response. Also, instead of answering
    in the first-person view of "we", answer in the third-person point of view of "NAI". 

    Relevant pieces of previous responses:
    {history}

    (You do not need to use these pieces of information if not relevant)

    Current conversation:
    Human: {input}
    AI:"""
    PROMPT = PromptTemplate(input_variables=["history", "input"], template=_DEFAULT_TEMPLATE)
    
    outputs = []

    # Input information
    document = input("Name of RFP document:")

    start_datetime = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    for i in range(len(final_sections)):
        # Initializing memory vector store
        embedding_size = 1536 # Dimensions of the OpenAIEmbeddings
        index = faiss.IndexFlatL2(embedding_size)
        embedding_fn = OpenAIEmbeddings(openai_api_key="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").embed_query
        vectorstore = FAISS(embedding_fn, index, InMemoryDocstore({}), {}) 

        # Create VectorStoreRetrieverMemory
        retriever = vectorstore.as_retriever(search_kwargs=dict(k=2))
        memory = VectorStoreRetrieverMemory(retriever=retriever, save_responses=False)

        custom_memory(memory)

        # Initialize Conversation Chain
        conversation = ConversationChain(llm = llm,
                                        prompt=PROMPT,
                                        memory = memory,
                                        verbose = True)
        
        #user_input = sections[i] + ' ' + contents[i]
        user_input = final_sections[i]

        # Input template
        template = """Based on your previous responses to similar sections, respond to the following section (and its contents) in a similar manner to previous, related responses, or if you have responded to this exact section before, return to me the identical previous response. Start the answer with "NAI agrees", "NAI understands", "NAI acknowledges", or "NAI will". Also, instead of answering in a first-person point of view as "we", answer in a third-person point of view as "NAI". Replace all instances of "we" with "NAI". Here is the section and it's contents: {section}"""
        prompt = PromptTemplate.from_template(template) 
        final_prompt = prompt.format(section=user_input)

        # AI generation
        predict = conversation.predict(input=final_prompt)
        print(predict)

        outputs.append(predict)

    for i in range(len(final_sections)):
        final_sections[i] = final_sections[i][0:25]
    
    finish_datetime = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    generate_doc(document, final_sections, outputs, start_datetime, finish_datetime)


def pdftodocx():
    if not exists('C:/Users/AChan/Documents/Source_Document_Store/3.a - FCC_SOW.docx'):
        pdf_file = 'C:/nai_localgpt/localGPT/SOURCE_DOCUMENTS/3.a - FCC_SOW.pdf'
        docx_file = 'C:/Users/AChan/Documents/Source_Document_Store/3.a - FCC_SOW.docx'
        cv = Converter(pdf_file)
        cv.convert(docx_file)    
        cv.close()


def extract_headers(doc):
    # Regex pattern
    pattern = '(\d+(?:\.\d+)+)\s+((?![A-Z][a-z])(?!^\d+$).+)'

    matches = re.findall(pattern, doc.page_content)
    index = 0
    result = []
    #print('\nExtracting headers:\n')
    #print(matches)
    #print('\nCleaning headers\n')
    for match in matches:
        tuple_to_list = list(match)
        tuple_to_list[1] = tuple_to_list[1].strip()
        tuple_to_list[1] = tuple_to_list[1].replace('.', '')
        list_to_tuple = tuple(tuple_to_list)
        if (list_to_tuple[1].isnumeric()):
            matches.pop(index)
            #print('popped!')
            index += 1
        else:
            print(list_to_tuple)
            result.append(list_to_tuple)
            index += 1
    return result

def match_and_extract(doc, start_pattern, stop_pattern):
    recording = False
    output_section = []
    doc_lines = doc.page_content.splitlines()
    stop_index = 0

    for line in doc_lines:
        if stop_index > 30:
            break
            
        line = re.sub('\t', '', line)
        if recording:
            #if re.search(stop_pattern, line) is not None:
            if stop_pattern in line:
                recording = False
                continue
            
            if line == '':
                continue

            output_section.append(line.strip())
            stop_index += 1
        #elif re.search(start_pattern, line) is not None:
        elif start_pattern in line:
            recording = True
            output_section = []
            output_section.append(line.strip())

        

    return output_section

def query_sql_connection(question, answer, time, username, source_doc):
    # Define the connection string
    connection_string = r'Driver={SQL Server};Server=DB01;Database=AIDB;UID=LocalGPTDBUser;PWD=<D?s5V?z\7dz3'

    # Establish the connection
    conn = pyodbc.connect(connection_string)

    try:
        # Create a cursor
        cursor = conn.cursor()
        # cursor.execute("SET SESSION TRANSACTION ISOLATION LEVEL READ COMMITTED")

        # Execute a SQL query
        # Define the insert statement
        insert_query = """
            INSERT INTO [dbo].[LocalGPTPromptLog] (
                [Question], [Answer], [JSONObj], [ProcessTimeSeconds], [TaskName], [LLMModelName], [EmbeddingsModelName], 
                [CreatedBy], [CreatedDateTime]
            ) OUTPUT INSERTED.LocalGPTPromptLogID
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """

        # Prepare the data for insertion
        question = question
        answer = answer
        json_obj = '{"key": "value"}'
        process_time_seconds = time
        created_by = username
        created_datetime = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  

        # Execute the insert statement
        cursor.execute(insert_query, question, answer, json_obj, process_time_seconds, TASK_NAME, LLM_MODEL, EMBEDDINGS_MODEL, created_by, created_datetime)
        referenced_id = cursor.fetchone()[0]

        # Commit the transaction
        conn.commit()

        # Retrieve the newly generated primary key value from the referenced table
        # Insert into the table with the foreign key constraint
        insert_query = """
            INSERT INTO LocalGPTPromptLogSourceDocs (LocalGPTPromptLogID, SourceDocName)
            VALUES (?, ?)
        """

        source_doc = source_doc
        for document in source_doc:
            cursor.execute(insert_query, referenced_id, document.metadata["source"])
            
        conn.commit()

    finally:
        # Close the cursor and connection
        cursor.close()
        conn.close()


def convo_sql_connection(input, output, isValid):
# Define the connection string
    connection_string = r'Driver={SQL Server};Server=DB01;Database=AIDB;UID=LocalGPTDBUser;PWD=<D?s5V?z\7dz3'

    # Establish the connection
    conn = pyodbc.connect(connection_string)

    try:
        # Create a cursor
        cursor = conn.cursor()
        # cursor.execute("SET SESSION TRANSACTION ISOLATION LEVEL READ COMMITTED")

        # Execute a SQL query
        # Define the insert statement
        insert_query = """
            INSERT INTO [dbo].[LocalGPTMemory] (
                [Input], [Output], [IsValid], [CreatedDateTime]
            ) OUTPUT INSERTED.LocalGPTMemoryID
            VALUES (?, ?, ?, ?)
        """

        # Execute the insert statement
        cursor.execute(insert_query, input, output, isValid, datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        referenced_id = cursor.fetchone()[0]

        # Commit the transaction
        conn.commit()

    finally:
        # Close the cursor and connection
        cursor.close()
        conn.close()

def save_to_spreadsheet(document_name, section, input, output):
    try:
        # Try to open the existing spreadsheet
        workbook = openpyxl.load_workbook(f'{document_name}.xlsx')
        sheet = workbook.active
    except FileNotFoundError:
        # If the spreadsheet doesn't exist, create a new one
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet['A1'] = document_name
        sheet['A2'] = 'Section'
        sheet['B2'] = 'Input'
        sheet['C2'] = 'Output'
    
    next_row = sheet.max_row + 1
    sheet[f'A{next_row}'] = section
    sheet[f'B{next_row}'] = input
    sheet[f'C{next_row}'] = output

    workbook.save(f'{document_name}.xlsx')

def generate_doc(document_name, sections, outputs, start_datetime, finish_datetime):
    # New document
    document = Document()

    # Header
    section = document.sections[0]
    header = section.header
    image = header.add_paragraph()
    image.add_run().add_picture('C:/nai_localgpt/localGPT/NAI_Logo_72dpi.png', width=Inches(1.9))
    image.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
    paragraph = header.add_paragraph('Response to ' + document_name)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Title page
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run('\nVOL - RFP Response\n').font.size = Pt(14) 
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(document_name + '\n').font.size = Pt(14) 
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run('Submitted To').font.size = Pt(12) 
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run('Contact:\nTelephone:').font.size = Pt(12) 
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run('Prepared By').font.size = Pt(12)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run('North Atlantic Industries, Inc. \n116 Wilbur Place\nBohemia, New York 11716').font.size = Pt(12)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run('DUNS # 79-346-0940\nCage Code: 0VGU1').font.size = Pt(12)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run('NAI Contact:\nRon Stack – Western Area Sales Manager\nTelephone: (951) 264-7027 / Email: rstack@naii.com').font.size = Pt(12)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run('Start Time: ' + start_datetime + '\nFinish Time: ' + finish_datetime).font.size = Pt(12)
    document.add_page_break()
    
    # Sections
    if len(sections) == len(outputs):
        for i in range(len(sections)):
            try:
                document.add_heading(sections[i], level=1)
                document.add_paragraph(outputs[i])
            except:
                document.save(document_name + '.docx')
        document.save(document_name + '.docx')
    else:
        for i in range(min(len(sections), len(outputs))):
            try:
                document.add_heading(sections[i], level=1)
                document.add_paragraph(outputs[i])
            except:
                document.save(document_name + '.docx')
        document.save(document_name + '.docx')

            
        
        
    
    


if __name__ == "__main__":
    main()  
